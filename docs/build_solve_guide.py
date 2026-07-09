from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = r"C:\Users\최준용\Documents\[이스트캠프] 워게임 중간 대회\docs\WARGAME_SOLVE_GUIDE.docx"
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')

def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(sum(widths)))
    tblW.set(qn('w:type'), 'dxa')
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement('w:tblGrid')
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(w))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(widths[idx]))
            tcW.set(qn('w:type'), 'dxa')
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)

def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Malgun Gothic'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [('Heading 1',16,BLUE,18,10),('Heading 2',13,BLUE,14,7),('Heading 3',12,DARK,10,5)]:
        st = styles[name]
        st.font.name = 'Malgun Gothic'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    code = styles.add_style('CodeBlock', 1)
    code.font.name = 'Consolas'
    code._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.1)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.1

def add_kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    set_table_widths(table, [1700, 7660])
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], 'E8EEF5')
        for c in cells:
            set_cell_margins(c)
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.name = 'Malgun Gothic'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
                    run.font.size = Pt(9.5)
        cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

def add_code(doc, text):
    p = doc.add_paragraph(style='CodeBlock')
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

def add_steps(doc, steps):
    for step in steps:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(4)
        p.add_run(step)

def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

doc = Document()
style_doc(doc)
p = doc.add_paragraph()
r = p.add_run('EST Wargame Lab 풀이 가이드')
r.font.name = 'Malgun Gothic'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
r.font.size = Pt(24)
r.font.bold = True
p.paragraph_format.space_after = Pt(3)
p = doc.add_paragraph('대상 문제: Session Shadow, Blind Notice, Image Vault, Admin Memo Chain')
p.runs[0].font.color.rgb = RGBColor(85, 85, 85)
doc.add_paragraph('이 문서는 먼저 공략을 따라 풀어보며 취약점 흐름을 익히는 용도입니다. 각 문제는 관찰 포인트, 핵심 원리, 단계별 풀이, 실패 시 확인할 부분을 포함합니다.')
add_kv_table(doc, [('사이트','http://100.83.178.9/ 또는 ISO 실행 환경의 로컬 IP'),('제출 방식','각 문제 페이지 하단의 플래그 제출 폼 사용'),('플래그 형식','EST{...}'),('주의','공격 코드는 이 과제용 워게임 VM 안에서만 사용하세요.')])

sections = [
('1. Session Shadow',[('난이도','중하'),('유형','Cookie 값 변조'),('목표','shadow_role 쿠키 값을 admin으로 변경해 플래그 확인'),('정답','EST{cookie_role_admin_shadow}')],'서버가 사용자의 역할을 세션 저장소가 아니라 브라우저 쿠키 shadow_role 값에서 그대로 읽습니다. 쿠키는 사용자가 직접 수정할 수 있으므로, 역할 검증을 쿠키 값만으로 처리하면 권한 상승이 발생합니다.',['문제 페이지에 접속한 뒤 개발자 도구의 Application 또는 Storage 탭을 엽니다.','Cookies 항목에서 현재 사이트의 shadow_role 값을 확인합니다. 처음 접속하면 guest로 설정되어 있습니다.','shadow_role 값을 admin으로 수정하고 저장합니다.','페이지를 새로고침하면 current role이 admin으로 바뀌고 플래그가 표시됩니다.','표시된 플래그를 문제 하단 제출 폼에 입력합니다.'],['shadow_role=guest','shadow_role=admin'],'shadow_role 쿠키가 보이지 않으면 문제 페이지를 한 번 새로고침하세요. 쿠키 값을 바꾼 뒤에도 guest로 보이면 도메인/path가 현재 사이트와 맞는 쿠키를 수정했는지 확인해야 합니다.'),
('2. Blind Notice',[('난이도','중'),('유형','Boolean-based Blind SQL Injection'),('목표','검색 결과의 존재/부재 차이로 challenge_flags의 flag 값을 한 글자씩 추출'),('정답','EST{boolean_blind_notice_42c7}')],'검색어가 SQL의 LIKE 문자열 안에 그대로 이어 붙습니다. 화면에는 데이터가 직접 출력되지 않지만, 조건이 참이면 공지 존재, 거짓이면 공지 없음으로 갈라집니다. 이 차이를 이용해 플래그의 각 글자 ASCII 값을 추측할 수 있습니다.',['일반 검색어 welcome을 넣어 정상 검색 결과가 나오는지 확인합니다.','작은따옴표 또는 주석을 포함한 입력으로 SQL 문맥을 깨뜨릴 수 있는지 확인합니다.','참 조건과 거짓 조건을 하나씩 넣어 화면 응답이 달라지는지 비교합니다.','SUBSTRING과 ASCII를 이용해 blind-notice 플래그의 첫 번째 글자가 특정 ASCII 값인지 검사합니다.','위 과정을 위치 1부터 반복해 전체 플래그를 복원합니다.'],["%' AND 1=1 -- ","%' AND 1=2 -- ","%' AND ASCII(SUBSTRING((SELECT flag FROM challenge_flags WHERE challenge_id='blind-notice'),1,1))=69 -- ",'위 payload에서 위치 1과 ASCII 값 69를 바꿔가며 EST{...} 전체를 추출합니다.'],'URL에서 공백과 특수문자가 깨지면 주석이 제대로 동작하지 않습니다. 브라우저 주소창에서 직접 입력할 때는 공백을 %20으로 인코딩하거나, Burp Suite/Reapter 같은 도구에서 파라미터를 보내는 편이 안정적입니다.'),
('3. Image Vault',[('난이도','중상'),('유형','File Path Filter Bypass'),('목표','.png 포함 여부만 검사하는 파일명 필터를 우회해 flag.txt 읽기'),('정답','EST{image_vault_path_filter_bypass}')],'서버는 요청한 file 값에 .png 문자열이 들어있는지만 확인합니다. 실제 확장자, 정규화된 경로, 허용 디렉터리 이탈 여부는 검증하지 않습니다. 따라서 allowed.png 같은 디렉터리명을 경유하고 ../를 사용하면 .png 검사를 통과하면서 flag.txt로 이동할 수 있습니다.',['기본 입력 cover.png를 열어 정상적으로 파일 내용이 출력되는지 확인합니다.','.png가 없는 flag.txt를 직접 요청하면 필터에 막히는 것을 확인합니다.','서버가 .png 문자열 포함 여부만 보는 점을 이용해 allowed.png/../flag.txt 형태의 경로를 만듭니다.','file 파라미터에 해당 값을 넣고 열기를 누르면 실제 경로가 assets/vault/flag.txt로 해석됩니다.','출력된 플래그를 페이지 하단 제출 폼에 넣습니다.'],['/challenges/image-vault.php?file=flag.txt','/challenges/image-vault.php?file=allowed.png/../flag.txt'],'cover.png/../flag.txt는 실패할 수 있습니다. cover.png는 파일이고 디렉터리가 아니기 때문입니다. 경로 중간에 들어가는 .png 요소는 디렉터리여야 하므로 allowed.png/../flag.txt를 사용해야 합니다.'),
('4. Admin Memo Chain',[('난이도','상'),('유형','Stored XSS Chain'),('목표','저장형 XSS 페이로드로 관리자 전용 내부 API 접근을 유도'),('정답','EST{stored_xss_admin_memo_chain}')],'메모 내용이 HTML 이스케이프 없이 다시 출력됩니다. 사용자가 작성한 script가 관리자 검토 흐름에서 실행된다고 가정하면, 일반 사용자는 직접 접근할 수 없는 admin_api=flag 엔드포인트를 관리자 컨텍스트에서 호출하게 만들 수 있습니다. 이 문제는 학습용으로 관리자 검토 버튼이 해당 흐름을 시뮬레이션합니다.',['메모에 일반 HTML 태그를 넣고 저장해 화면에 그대로 반영되는지 확인합니다.','script 태그가 필터링되지 않는다는 점을 확인합니다.','관리자 전용 API 경로 /challenges/admin-memo-chain.php?admin_api=flag를 호출하는 fetch 페이로드를 작성합니다.','메모를 저장한 뒤 관리자 검토 요청 버튼을 누릅니다.','최근 메모의 review_result 영역에 admin browser fetched private API 메시지와 함께 플래그가 기록됩니다.'],["<script>fetch('/challenges/admin-memo-chain.php?admin_api=flag')</script>",'검토 결과: admin browser fetched private API: EST{stored_xss_admin_memo_chain}'],'fetch 경로 문자열이 다르면 학습용 관리자 검토 시뮬레이터가 인식하지 못합니다. 작은따옴표 또는 큰따옴표 형태의 정확한 경로를 사용하세요.')
]

for title, meta, principle, steps, codes, fail in sections:
    page_break(doc)
    doc.add_heading(title, level=1)
    add_kv_table(doc, meta)
    doc.add_heading('핵심 원리', level=2)
    doc.add_paragraph(principle)
    doc.add_heading('단계별 풀이', level=2)
    add_steps(doc, steps)
    doc.add_heading('페이로드 / 명령 예시', level=2)
    for code in codes:
        add_code(doc, code)
    doc.add_heading('실패 시 확인할 부분', level=2)
    doc.add_paragraph(fail)

doc.save(OUT)
print(OUT)