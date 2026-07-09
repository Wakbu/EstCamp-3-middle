from docx import Document
p=r'C:\Users\최준용\Documents\[이스트캠프] 워게임 중간 대회\docs\WARGAME_SOLVE_GUIDE.docx'
d=Document(p)
for i,x in enumerate(d.paragraphs[:25]):
    print(i, x.text.encode('unicode_escape').decode())
for ti,t in enumerate(d.tables[:2]):
    print('TABLE',ti)
    for row in t.rows:
        print([c.text.encode('unicode_escape').decode() for c in row.cells])