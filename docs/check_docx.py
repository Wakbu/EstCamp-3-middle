from docx import Document
p=r'C:\Users\최준용\Documents\[이스트캠프] 워게임 중간 대회\docs\WARGAME_SOLVE_GUIDE.docx'
d=Document(p)
print('paragraphs', len(d.paragraphs))
print('tables', len(d.tables))
print('first', d.paragraphs[0].text)
print('headings', [x.text for x in d.paragraphs if x.style.name.startswith('Heading 1')])